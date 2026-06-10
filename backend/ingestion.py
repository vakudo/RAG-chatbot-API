import csv
import json
import threading
from pathlib import Path

import fitz
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_text_splitters import (
    MarkdownHeaderTextSplitter,
    RecursiveCharacterTextSplitter,
)
from openpyxl import load_workbook

from backend.config import settings
from backend.db import connect
from backend.retriever import get_vectorstore

_index_lock = threading.Lock()

ALLOWED_EXTENSIONS = {
    ".pdf", ".txt", ".md", ".csv", ".docx", ".xlsx", ".xlsm", ".html", ".htm",
}


def _index_path() -> Path:
    return Path(settings.uploads_path) / "index.json"


def load_index() -> dict:
    path = _index_path()
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return {}


def _save_index(index: dict) -> None:
    _index_path().write_text(json.dumps(index, indent=2), encoding="utf-8")


def _extract_xlsx(path: Path) -> str:
    # Rows become "header: value" lines so each chunk stays self-describing
    # after splitting, which keeps similarity search meaningful.
    wb = load_workbook(path, read_only=True, data_only=True)
    try:
        sheets = []
        for ws in wb.worksheets:
            header = None
            lines = [f"Sheet: {ws.title}"]
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c).strip() for c in row]
                if not any(cells):
                    continue
                if header is None:
                    header = cells
                    continue
                pairs = [f"{h}: {v}" for h, v in zip(header, cells) if v and h]
                if pairs:
                    lines.append("; ".join(pairs))
            if len(lines) > 1:
                sheets.append("\n".join(lines))
        return "\n\n".join(sheets)
    finally:
        wb.close()


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_csv(path: Path) -> str:
    # Same "header: value" flattening as Excel sheets.
    with open(path, newline="", encoding="utf-8", errors="ignore") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            dialect = csv.excel
        rows = [r for r in csv.reader(f, dialect) if any(c.strip() for c in r)]
    if not rows:
        return ""
    header, *data = rows
    if not data:
        return " | ".join(header)
    lines = []
    for row in data:
        pairs = [f"{h}: {v}" for h, v in zip(header, row) if v.strip() and h.strip()]
        if pairs:
            lines.append("; ".join(pairs))
    return "\n".join(lines)


def _extract_html(path: Path) -> str:
    soup = BeautifulSoup(
        path.read_text(encoding="utf-8", errors="ignore"), "html.parser"
    )
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _ocr_pdf(path: Path) -> str:
    """Fallback for scanned PDFs with no text layer."""
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        return ""
    ocr = RapidOCR()
    pages = []
    with fitz.open(path) as doc:
        for page in doc:
            png = page.get_pixmap(dpi=200).tobytes("png")
            result, _ = ocr(png)
            if result:
                pages.append("\n".join(line[1] for line in result))
    return "\n".join(pages)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        with fitz.open(path) as doc:
            text = "\n".join(page.get_text() for page in doc)
        if len(text.strip()) < 50:
            ocr_text = _ocr_pdf(path)
            if ocr_text.strip():
                return ocr_text
        return text
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix in {".html", ".htm"}:
        return _extract_html(path)
    # .txt, .md and anything else readable as plain text
    return path.read_text(encoding="utf-8", errors="ignore")


def split_text(text: str, suffix: str) -> list[str]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap
    )
    if suffix != ".md":
        return splitter.split_text(text)
    # Markdown: split along headings first so chunks follow the document
    # structure; each chunk is prefixed with its heading path for context.
    header_splitter = MarkdownHeaderTextSplitter(
        [("#", "h1"), ("##", "h2"), ("###", "h3")]
    )
    chunks = []
    for section in header_splitter.split_text(text):
        prefix = " / ".join(section.metadata.values())
        for piece in splitter.split_text(section.page_content):
            chunks.append(f"{prefix}\n{piece}" if prefix else piece)
    return chunks or splitter.split_text(text)


def ingest_file(path: Path, doc_id: str, filename: str) -> int:
    text = extract_text(path)
    chunks = split_text(text, path.suffix.lower())
    if not chunks:
        raise ValueError("No text could be extracted from the file")

    metadatas = [
        {"doc_id": doc_id, "filename": filename, "chunk_index": i}
        for i in range(len(chunks))
    ]
    get_vectorstore().add_texts(chunks, metadatas=metadatas)

    with _index_lock:
        index = load_index()
        index[doc_id] = {"filename": filename, "chunks_count": len(chunks)}
        _save_index(index)
    return len(chunks)


def remove_document(doc_id: str) -> None:
    with connect() as conn:
        conn.execute(
            "DELETE FROM langchain_pg_embedding WHERE cmetadata->>'doc_id' = %s",
            (doc_id,),
        )
    for path in Path(settings.uploads_path).glob(f"{doc_id}_*"):
        path.unlink(missing_ok=True)
    with _index_lock:
        index = load_index()
        index.pop(doc_id, None)
        _save_index(index)
